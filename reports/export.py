import io
from datetime import date

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Table, TableStyle

BRAND = "#4f46e5"
INK = "#1e293b"
MUTED = "#64748b"
LINE = "#e5e9f2"

SESSION_COLS = [
    "Date", "Sess", "Class", "Subject", "Topic Taught", "Trainer",
    "Location", "Per", "Start", "End",
]


def _fmt_time(value):
    return value.strftime("%H:%M") if value else "-"


def _session_table_rows(sessions):
    rows = []
    for s in sessions:
        rows.append([
            s.date.strftime("%d %b %Y"),
            str(s.session_number or "-"),
            str(s.school_class),
            s.subject.name,
            s.topic_taught,
            s.trainer.full_name,
            s.location or "-",
            str(s.period or "-"),
            _fmt_time(s.start_time),
            _fmt_time(s.end_time),
        ])
    return rows


def _pdf_table(data, col_widths, font_size=9):
    table = Table(data, colWidths=col_widths, repeatRows=1)
    style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(BRAND)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), font_size),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor(LINE)),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]
    table.setStyle(TableStyle(style))
    return table


def build_pdf(context):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=landscape(A4),
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
        title="OneFuture Organization Report",
    )

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle(
        "ReportTitle", parent=styles["Title"], fontSize=20,
        textColor=colors.HexColor(INK), spaceAfter=2,
    )
    meta = ParagraphStyle(
        "Meta", parent=styles["Normal"], fontSize=9,
        textColor=colors.HexColor(MUTED), spaceAfter=10,
    )
    h2 = ParagraphStyle(
        "ReportH2", parent=styles["Heading2"], fontSize=13,
        textColor=colors.HexColor(BRAND), spaceBefore=14, spaceAfter=6,
    )

    story = []
    story.append(Paragraph("OneFuture — Organization Report", h1))
    story.append(Paragraph(
        f"Generated on {date.today().strftime('%d %B %Y')} &nbsp;·&nbsp; "
        f"Covering {context['total_sessions']} sessions", meta,
    ))

    story.append(Paragraph("Filters", h2))
    story.append(_pdf_table(
        [["Filter", "Value"]] + [[k, v] for k, v in context["filter_summary"]],
        [70 * mm, 180 * mm],
    ))

    story.append(Paragraph("Summary", h2))
    story.append(_pdf_table(
        [["Total Sessions", "Active Trainers", "Classes Covered", "Subjects Covered"],
         [
             str(context["total_sessions"]),
             str(context["active_trainers"]),
             str(context["classes_count"]),
             str(context["subjects_count"]),
         ]],
        [60 * mm, 60 * mm, 60 * mm, 70 * mm],
    ))

    story.append(Paragraph("Sessions by Trainer", h2))
    story.append(_pdf_table(
        [["Trainer", "Sessions"]]
        + [[r["trainer__full_name"], str(r["count"])] for r in context["by_trainer"]],
        [200 * mm, 50 * mm],
    ))

    story.append(Paragraph("Sessions by Class", h2))
    story.append(_pdf_table(
        [["Class", "Sessions"]]
        + [[f'{r["school_class__name"]}{" - " + r["school_class__section"] if r["school_class__section"] else ""}', str(r["count"])] for r in context["by_class"]],
        [200 * mm, 50 * mm],
    ))

    story.append(Paragraph("Sessions by Subject", h2))
    story.append(_pdf_table(
        [["Subject", "Sessions"]]
        + [[r["subject__name"], str(r["count"])] for r in context["by_subject"]],
        [200 * mm, 50 * mm],
    ))

    story.append(Paragraph("Session Details", h2))
    if context["sessions"]:
        story.append(_pdf_table(
            [SESSION_COLS] + _session_table_rows(context["sessions"]),
            [26 * mm, 11 * mm, 30 * mm, 30 * mm, 62 * mm, 38 * mm, 28 * mm, 10 * mm, 15 * mm, 15 * mm],
            font_size=8,
        ))
    else:
        story.append(Paragraph("No sessions match the selected filters.", styles["Normal"]))

    doc.build(story)
    return buf.getvalue()


def _shade_cell(cell, hex_color):
    fill = OxmlElement("w:shd")
    fill.set(qn("w:val"), "clear")
    fill.set(qn("w:fill"), hex_color.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(fill)


def _docx_pair_table(doc, data):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(value))
            run.font.size = Pt(9)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string("ffffff")
                _shade_cell(cell, BRAND)


def build_docx(context):
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    doc.add_heading("OneFuture — Organization Report", 0)
    meta = doc.add_paragraph()
    run = meta.add_run(
        f"Generated on {date.today().strftime('%d %B %Y')}  |  "
        f"Covering {context['total_sessions']} sessions"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED.lstrip("#"))

    doc.add_heading("Filters", level=1)
    _docx_pair_table(doc, [["Filter", "Value"]] + [[k, v] for k, v in context["filter_summary"]])

    doc.add_heading("Summary", level=1)
    _docx_pair_table(doc, [
        ["Total Sessions", "Active Trainers", "Classes Covered", "Subjects Covered"],
        [str(context["total_sessions"]), str(context["active_trainers"]),
         str(context["classes_count"]), str(context["subjects_count"])],
    ])

    doc.add_heading("Sessions by Trainer", level=1)
    _docx_pair_table(doc, [["Trainer", "Sessions"]]
                     + [[r["trainer__full_name"], str(r["count"])] for r in context["by_trainer"]])

    doc.add_heading("Sessions by Class", level=1)
    _docx_pair_table(doc, [["Class", "Sessions"]]
                     + [[f'{r["school_class__name"]}{" - " + r["school_class__section"] if r["school_class__section"] else ""}', str(r["count"])] for r in context["by_class"]])

    doc.add_heading("Sessions by Subject", level=1)
    _docx_pair_table(doc, [["Subject", "Sessions"]]
                     + [[r["subject__name"], str(r["count"])] for r in context["by_subject"]])

    doc.add_heading("Session Details", level=1)
    if context["sessions"]:
        data = [SESSION_COLS] + _session_table_rows(context["sessions"])
        table = doc.add_table(rows=len(data), cols=len(SESSION_COLS))
        table.style = "Table Grid"
        table.autofit = True
        for r_idx, row in enumerate(data):
            for c_idx, value in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(value))
                run.font.size = Pt(8)
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string("ffffff")
                    _shade_cell(cell, BRAND)
    else:
        doc.add_paragraph("No sessions match the selected filters.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()